"""MCP tool for generating employee CVs in DOCX format."""

from __future__ import annotations

import logging
import os
import uuid
from pathlib import Path
from typing import Annotated

from fastmcp import Context

from .app import mcp, _pg
from .pg_age_helper import _sanitize_sql_string

logger = logging.getLogger("talent_mcp.cv")

# Output directory for generated CVs
CV_OUTPUT_DIR = Path(os.getenv("CV_OUTPUT_DIR", "/tmp/talentiq_cvs"))
CV_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Base URL for downloads (set by backend, defaults to localhost)
CV_DOWNLOAD_BASE = os.getenv("CV_DOWNLOAD_BASE", "http://localhost:8000/af/cv/files")


@mcp.tool
async def list_cv_templates(
    ctx: Context = None,
) -> list[dict]:
    """List available CV templates. The user must choose one before generating a CV."""
    if ctx:
        await ctx.info("[QUERY] CV: Listing available templates")

    template_dir = Path(__file__).resolve().parents[1] / "agent" / "template_docs"
    templates = []
    if template_dir.is_dir():
        for f in sorted(template_dir.iterdir()):
            if f.suffix.lower() in (".docx", ".pdf"):
                fmt = f.suffix.lstrip(".").lower()
                templates.append({
                    "id": f.stem,
                    "name": f.stem.replace("_", " ").replace("-", " "),
                    "filename": f.name,
                    "format": fmt,
                    "usable": fmt == "docx",
                    "note": "DOCX — can generate CV" if fmt == "docx" else "PDF — preview/reference only, cannot generate from this template",
                })
    return templates


@mcp.tool
async def generate_employee_cv(
    employee_email: Annotated[str, "Employee email address (e.g., 'jessica.berry@dxc.com')"],
    graph_name: Annotated[str, "Graph name for AGE cypher calls"],
    template_name: Annotated[str, "Template filename from template_docs (e.g., '01 CV_Coordinador.docx'). Empty for default."] = "",
    format: Annotated[str, "Output format: 'docx' (default)"] = "docx",
    anonymize: Annotated[bool, "If true, remove personal identifiers from CV"] = False,
    ctx: Context = None,
) -> dict:
    """Generate a standardized CV/resume for an employee in DOCX format.

    Queries the talent graph for all employee data and generates a professional
    DXC-branded CV document.
    """
    if ctx:
        await ctx.info(f"[cv] Generating CV for {employee_email}...")

    logger.info("[cv] Generating CV for %s (format=%s, anonymize=%s)", employee_email, format, anonymize)

    # Query employee profile
    email_safe = _sanitize_sql_string(employee_email)

    profile_sql = f"""SELECT * FROM ag_catalog.cypher('{_sanitize_sql_string(graph_name)}', $$
        MATCH (e:Employee) WHERE e.email = '{email_safe}'
        RETURN e.name AS name, e.first_name AS first_name, e.last_name AS last_name,
               e.email AS email, e.phone AS phone, e.job_title AS job_title,
               e.job_level AS job_level, e.skill_level AS skill_level,
               e.years_of_experience AS yoe, e.education_degree AS degree,
               e.education_field AS field, e.resume_summary AS summary,
               e.delivery_model AS delivery_model, e.eqf_level AS eqf_level
    $$) AS (name ag_catalog.agtype, first_name ag_catalog.agtype, last_name ag_catalog.agtype,
            email ag_catalog.agtype, phone ag_catalog.agtype, job_title ag_catalog.agtype,
            job_level ag_catalog.agtype, skill_level ag_catalog.agtype,
            yoe ag_catalog.agtype, degree ag_catalog.agtype,
            field ag_catalog.agtype, summary ag_catalog.agtype,
            delivery_model ag_catalog.agtype, eqf_level ag_catalog.agtype);"""

    profile_rows = await _pg().query_using_sql_cypher(profile_sql, graph_name)
    if not profile_rows:
        return {"error": f"Employee not found: {employee_email}"}

    profile = profile_rows[0]

    # Query skills
    skills_sql = f"""SELECT * FROM ag_catalog.cypher('{_sanitize_sql_string(graph_name)}', $$
        MATCH (e:Employee)-[hs:HAS_SKILL]->(s:Skill)
        WHERE e.email = '{email_safe}'
        RETURN s.name AS skill, hs.level AS level, hs.years_of_experience AS yoe, hs.is_primary AS is_primary
    $$) AS (skill ag_catalog.agtype, level ag_catalog.agtype, yoe ag_catalog.agtype, is_primary ag_catalog.agtype);"""

    skills = await _pg().query_using_sql_cypher(skills_sql, graph_name)

    # Query certifications
    certs_sql = f"""SELECT * FROM ag_catalog.cypher('{_sanitize_sql_string(graph_name)}', $$
        MATCH (e:Employee)-[hc:HOLDS_CERT]->(c:Certification)
        WHERE e.email = '{email_safe}'
        RETURN c.name AS cert, hc.status AS status, hc.issue_date AS issued, hc.expiry_date AS expiry
    $$) AS (cert ag_catalog.agtype, status ag_catalog.agtype, issued ag_catalog.agtype, expiry ag_catalog.agtype);"""

    certs = await _pg().query_using_sql_cypher(certs_sql, graph_name)

    # Query languages
    langs_sql = f"""SELECT * FROM ag_catalog.cypher('{_sanitize_sql_string(graph_name)}', $$
        MATCH (e:Employee)-[sp:SPEAKS]->(l:Language)
        WHERE e.email = '{email_safe}'
        RETURN l.name AS language, sp.level AS level, sp.is_native AS is_native
    $$) AS (language ag_catalog.agtype, level ag_catalog.agtype, is_native ag_catalog.agtype);"""

    langs = await _pg().query_using_sql_cypher(langs_sql, graph_name)

    # Query education
    edu_sql = f"""SELECT * FROM ag_catalog.cypher('{_sanitize_sql_string(graph_name)}', $$
        MATCH (e:Employee)-[sa:STUDIED_AT]->(u:University)
        WHERE e.email = '{email_safe}'
        RETURN u.name AS university, sa.degree AS degree, sa.field AS field, sa.graduation_year AS year
    $$) AS (university ag_catalog.agtype, degree ag_catalog.agtype, field ag_catalog.agtype, year ag_catalog.agtype);"""

    edu = await _pg().query_using_sql_cypher(edu_sql, graph_name)

    # Query work experience (clients + projects)
    exp_sql = f"""SELECT * FROM ag_catalog.cypher('{_sanitize_sql_string(graph_name)}', $$
        MATCH (e:Employee)-[wf:WORKED_FOR]->(c:Client)
        WHERE e.email = '{email_safe}'
        RETURN c.name AS client, wf.role AS role, wf.project AS project,
               wf.start_date AS start_date, wf.end_date AS end_date, wf.is_current AS is_current
    $$) AS (client ag_catalog.agtype, role ag_catalog.agtype, project ag_catalog.agtype,
            start_date ag_catalog.agtype, end_date ag_catalog.agtype, is_current ag_catalog.agtype);"""

    exp = await _pg().query_using_sql_cypher(exp_sql, graph_name)

    # Query location
    loc_sql = f"""SELECT * FROM ag_catalog.cypher('{_sanitize_sql_string(graph_name)}', $$
        MATCH (e:Employee)-[:LOCATED_IN]->(l:Location)-[:IN_COUNTRY]->(c:Country)
        WHERE e.email = '{email_safe}'
        RETURN l.city AS city, c.name AS country
    $$) AS (city ag_catalog.agtype, country ag_catalog.agtype);"""

    loc = await _pg().query_using_sql_cypher(loc_sql, graph_name)

    if ctx:
        await ctx.info(f"[cv] Data collected: {len(skills)} skills, {len(certs)} certs, {len(langs)} languages, {len(edu)} education, {len(exp)} experience")

    # Resolve template — fail loudly on unknown names so the agent cannot
    # fabricate a template name and silently fall through to the default.
    TEMPLATE_DIR = Path(__file__).resolve().parents[1] / "agent" / "template_docs"
    available_docx = [f.name for f in TEMPLATE_DIR.iterdir() if f.suffix.lower() == ".docx"] if TEMPLATE_DIR.is_dir() else []
    template_path = None
    template_used = "default_dxc"  # tracks what we actually rendered with
    if template_name:
        candidate = TEMPLATE_DIR / template_name
        if candidate.exists() and candidate.suffix.lower() == ".docx":
            template_path = candidate
            template_used = template_name
            logger.info("[cv] Using template: %s", template_name)
        elif candidate.exists() and candidate.suffix.lower() == ".pdf":
            logger.warning("[cv] PDF template requested: %s", template_name)
            return {
                "error": f"Template '{template_name}' is a PDF and cannot be used for CV generation. PDF templates are for preview/reference only. Please choose a DOCX template instead.",
                "available_docx_templates": available_docx,
            }
        else:
            logger.warning("[cv] Template not found: %s (available: %s)", template_name, available_docx)
            return {
                "error": (
                    f"Template '{template_name}' does not exist. Only the templates listed in "
                    f"`available_docx_templates` are real. Do not invent template names. "
                    f"Call `list_cv_templates` and ask the user to pick one of the real options."
                ),
                "available_docx_templates": available_docx,
            }

    # Generate DOCX
    file_id = uuid.uuid4().hex[:12]

    # Strip agtype quotes from values
    def _clean(val):
        if val is None:
            return ""
        s = str(val).strip('"')
        return s if s != "None" and s != "null" else ""

    employee_name = _clean(profile.get("name"))
    if anonymize:
        employee_name = "Candidate Profile"

    filename = f"CV_{employee_name.replace(' ', '_')}_{file_id}.docx"
    filepath = CV_OUTPUT_DIR / filename

    _build_docx(
        filepath=filepath,
        profile=profile,
        skills=skills,
        certs=certs,
        langs=langs,
        edu=edu,
        exp=exp,
        loc=loc,
        anonymize=anonymize,
        clean=_clean,
        template_path=template_path,
    )

    download_url = f"{CV_DOWNLOAD_BASE}/{filename}"

    logger.info("[cv] CV generated: %s (%d bytes)", filepath, filepath.stat().st_size)
    if ctx:
        await ctx.info(f"[cv] CV generated: {filename}")

    return {
        "filename": filename,
        "download_url": download_url,
        "employee": employee_name,
        "template_used": template_used,
        "skills_count": len(skills),
        "certs_count": len(certs),
        "format": format,
    }


def _build_docx(filepath, profile, skills, certs, langs, edu, exp, loc, anonymize, clean, template_path=None):
    """Build a professional DOCX CV."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if template_path:
        # ── Template-preserving mode ──────────────────────────
        _build_docx_from_template(filepath, template_path, profile, skills, certs, langs, edu, exp, loc, anonymize, clean)
        return

    # ── Default mode (no template) ────────────────────────────
    doc = Document()

    def _has_style(name: str) -> bool:
        try:
            _ = doc.styles[name]
            return True
        except KeyError:
            return False

    def _add_heading(text: str, level: int = 2):
        style_name = f'Heading {level}'
        if _has_style(style_name):
            return doc.add_heading(text, level=level)
        else:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.bold = True
            run.font.size = Pt(14 if level == 1 else 12)
            run.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)
            return p

    def _add_bullet(text: str):
        if _has_style('List Bullet'):
            doc.add_paragraph(text, style='List Bullet')
        else:
            doc.add_paragraph(f"• {text}")

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # Header — DXC branding
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("DXC Technology")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)  # teal accent
    run.font.bold = True

    sub = header_para.add_run("\nProfessional Profile")
    sub.font.size = Pt(11)
    sub.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph("")  # spacer

    # Name & Title
    name = clean(profile.get("name")) if not anonymize else "Candidate Profile"
    title = clean(profile.get("job_title"))
    level = clean(profile.get("skill_level"))
    yoe = clean(profile.get("yoe"))

    name_para = _add_heading(name, level=1)
    if name_para.runs:
        name_para.runs[0].font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    subtitle = f"{title}"
    if level:
        subtitle += f" | {level}"
    if yoe:
        subtitle += f" | {yoe} years experience"
    doc.add_paragraph(subtitle).runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Contact & Location
    if not anonymize:
        contact_parts = []
        email_val = clean(profile.get("email"))
        phone_val = clean(profile.get("phone"))
        if email_val:
            contact_parts.append(email_val)
        if phone_val:
            contact_parts.append(phone_val)
        if loc:
            city = clean(loc[0].get("city"))
            country = clean(loc[0].get("country"))
            if city and country:
                contact_parts.append(f"{city}, {country}")
        if contact_parts:
            doc.add_paragraph(" | ".join(contact_parts))

    # Summary
    summary = clean(profile.get("summary"))
    if summary:
        _add_heading("Professional Summary", level=2)
        doc.add_paragraph(summary)

    # Skills
    if skills:
        _add_heading("Technical Skills", level=2)
        # Group by level
        skill_groups = {}
        for s in skills:
            lvl = clean(s.get("level")) or "Other"
            skill_name = clean(s.get("skill"))
            if lvl not in skill_groups:
                skill_groups[lvl] = []
            skill_groups[lvl].append(skill_name)

        for lvl in ["Expert", "Guru", "Advanced", "Intermediate", "Basic", "Other"]:
            if lvl in skill_groups:
                p = doc.add_paragraph()
                run = p.add_run(f"{lvl}: ")
                run.font.bold = True
                p.add_run(", ".join(skill_groups[lvl]))

    # Certifications
    if certs:
        _add_heading("Certifications", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "Certification"
        hdr[1].text = "Status"
        hdr[2].text = "Expiry"
        for c in certs:
            row = table.add_row().cells
            row[0].text = clean(c.get("cert"))
            row[1].text = clean(c.get("status"))
            row[2].text = clean(c.get("expiry"))

    # Education
    if edu:
        _add_heading("Education", level=2)
        for e in edu:
            uni = clean(e.get("university"))
            deg = clean(e.get("degree"))
            fld = clean(e.get("field"))
            yr = clean(e.get("year"))
            line = f"{deg}"
            if fld:
                line += f" in {fld}"
            if uni:
                line += f" — {uni}"
            if yr:
                line += f" ({yr})"
            _add_bullet(line)

    # Languages
    if langs:
        _add_heading("Languages", level=2)
        lang_parts = []
        for l in langs:
            lang_name = clean(l.get("language"))
            lvl = clean(l.get("level"))
            native = l.get("is_native")
            if native and str(native).lower() in ("true", "1"):
                lang_parts.append(f"{lang_name} (Native)")
            elif lvl:
                lang_parts.append(f"{lang_name} ({lvl})")
            else:
                lang_parts.append(lang_name)
        doc.add_paragraph(", ".join(lang_parts))

    # Work Experience
    if exp:
        _add_heading("Professional Experience", level=2)
        # Sort: current first, then by start_date desc
        sorted_exp = sorted(exp, key=lambda x: (
            str(x.get("is_current", "")).lower() not in ("true", "1"),
            clean(x.get("start_date")) or "0000",
        ))
        for e in sorted_exp:
            client_name = clean(e.get("client"))
            if anonymize:
                client_name = "Confidential Client"
            role = clean(e.get("role"))
            project = clean(e.get("project"))
            start = clean(e.get("start_date"))
            end = clean(e.get("end_date")) or "Present"
            is_current = str(e.get("is_current", "")).lower() in ("true", "1")

            p = doc.add_paragraph()
            run = p.add_run(f"{role}")
            run.font.bold = True
            p.add_run(f" at {client_name}")
            if start:
                period = f" ({start} — {'Present' if is_current else end})"
                p.add_run(period).font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            if project:
                _add_bullet(f"Project: {project}")

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated by TalentIQ — DXC Technology")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    doc.save(str(filepath))


def _build_docx_from_template(filepath, template_path, profile, skills, certs, langs, edu, exp, loc, anonymize, clean):
    """Fill a DXC DOCX template preserving its layout (textboxes, styles, headers/footers).

    The template uses:
    - Textboxes in the sidebar for: Formación (education), Idiomas (languages), Informática (skills)
    - Body paragraphs for: name (Bio Name style), summary (DXC Body Text), experience (Date company / Body Bullet)
    - Custom DXC styles: Bio Name, DXC Body Text, Date company, Body Bullet, DXC Section Title, etc.
    """
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree
    import copy

    doc = Document(str(template_path))
    body = doc.element.body
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # ── Helper: clear paragraphs in an XML container ─────────
    def _clear_paragraphs(container):
        """Remove all w:p elements from a container."""
        for p in list(container.findall(qn('w:p'))):
            container.remove(p)

    def _add_para(container, text, style_id=None, bold=False):
        """Add a paragraph to an XML container."""
        p = etree.SubElement(container, qn('w:p'))
        if style_id:
            pPr = etree.SubElement(p, qn('w:pPr'))
            pStyle = etree.SubElement(pPr, qn('w:pStyle'))
            pStyle.set(qn('w:val'), style_id)
        r = etree.SubElement(p, qn('w:r'))
        if bold:
            rPr = etree.SubElement(r, qn('w:rPr'))
            b_el = etree.SubElement(rPr, qn('w:b'))
        t = etree.SubElement(r, qn('w:t'))
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        return p

    # ── 1. Fill sidebar textboxes ────────────────────────────
    # Find the FIRST txbxContent (the main one — the second is a VML fallback copy)
    TXBX_TAG = '{%s}txbxContent' % W
    textboxes = [el for el in body.iter() if el.tag == TXBX_TAG]
    if textboxes:
        sidebar = textboxes[0]
        _clear_paragraphs(sidebar)

        # Education section
        _add_para(sidebar, "Education", bold=True)
        if edu:
            for e in edu:
                deg = clean(e.get("degree"))
                fld = clean(e.get("field"))
                uni = clean(e.get("university"))
                yr = clean(e.get("year"))
                line = f"{deg}"
                if fld:
                    line += f" in {fld}"
                if uni:
                    line += f" — {uni}"
                if yr:
                    line += f" ({yr})"
                _add_para(sidebar, line)
        else:
            _add_para(sidebar, "Not available")

        _add_para(sidebar, "")  # spacer

        # Languages section
        _add_para(sidebar, "Languages", bold=True)
        if langs:
            for l in langs:
                lang_name = clean(l.get("language"))
                lvl = clean(l.get("level"))
                native = l.get("is_native")
                if native and str(native).lower() in ("true", "1"):
                    _add_para(sidebar, f"{lang_name} (Native)")
                elif lvl:
                    _add_para(sidebar, f"{lang_name} ({lvl})")
                else:
                    _add_para(sidebar, lang_name)
        else:
            _add_para(sidebar, "Not available")

        _add_para(sidebar, "")  # spacer

        # Skills / Certifications section
        _add_para(sidebar, "Skills & Certifications", bold=True)
        if skills:
            skill_names = [clean(s.get("skill")) for s in skills if clean(s.get("skill"))]
            _add_para(sidebar, ", ".join(skill_names))
        if certs:
            cert_names = [clean(c.get("cert")) for c in certs if clean(c.get("cert"))]
            if cert_names:
                _add_para(sidebar, "")
                _add_para(sidebar, "Certifications:", bold=True)
                for cn in cert_names:
                    _add_para(sidebar, f"• {cn}")

        # Update VML fallback textbox (second txbxContent) with same content
        if len(textboxes) > 1:
            fallback = textboxes[1]
            _clear_paragraphs(fallback)
            for p in sidebar.findall(qn('w:p')):
                fallback.append(copy.deepcopy(p))

    # ── 2. Fill body paragraphs (right column) ───────────────
    # Remove existing body paragraphs but keep drawings/shapes (textboxes) and sectPr
    WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    VML_NS = "urn:schemas-microsoft-com:vml"
    for element in list(body):
        tag = element.tag
        # Keep sectPr
        if tag == qn('w:sectPr'):
            continue
        # Keep elements containing drawings (textbox shapes)
        has_drawing = element.find('.//{%s}drawing' % W) is not None
        has_pict = element.find('.//{%s}shape' % VML_NS) is not None
        has_wps = element.find('.//{%s}wsp' % WPS_NS) is not None
        if has_drawing or has_pict or has_wps:
            continue
        body.remove(element)

    # Find insertion point — before sectPr
    sect_pr = body.find(qn('w:sectPr'))
    insert_before = sect_pr

    def _insert_para(text, style_id=None, bold=False):
        """Insert a paragraph before sectPr."""
        p = etree.Element(qn('w:p'))
        if style_id:
            pPr = etree.SubElement(p, qn('w:pPr'))
            pStyle = etree.SubElement(pPr, qn('w:pStyle'))
            pStyle.set(qn('w:val'), style_id)
        r = etree.SubElement(p, qn('w:r'))
        if bold:
            rPr = etree.SubElement(r, qn('w:rPr'))
            b_el = etree.SubElement(rPr, qn('w:b'))
        t = etree.SubElement(r, qn('w:t'))
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        if insert_before is not None:
            body.insert(list(body).index(insert_before), p)
        else:
            body.append(p)
        return p

    # Name — use Bio Name style if available
    name = clean(profile.get("name")) if not anonymize else "Candidate Profile"
    _insert_para(name, style_id="BioName")

    # Summary
    summary = clean(profile.get("summary"))
    if summary:
        _insert_para(summary, style_id="DXCBodyText")

    # Principal Experience heading
    _insert_para("Professional Experience", style_id="DXCBodyText", bold=True)

    # Work Experience
    if exp:
        sorted_exp = sorted(exp, key=lambda x: (
            str(x.get("is_current", "")).lower() not in ("true", "1"),
            clean(x.get("start_date")) or "0000",
        ))
        for e in sorted_exp:
            client_name = clean(e.get("client"))
            if anonymize:
                client_name = "Confidential Client"
            role = clean(e.get("role"))
            project = clean(e.get("project"))
            start = clean(e.get("start_date"))
            end = clean(e.get("end_date")) or "Present"
            is_current = str(e.get("is_current", "")).lower() in ("true", "1")

            period_end = "Present" if is_current else end
            header = f"{start} — {period_end}. {role} at {client_name}"
            _insert_para(header, style_id="Datecompany")

            if project:
                _insert_para(f"Project: {project}", style_id="BodyBullet")

    doc.save(str(filepath))
